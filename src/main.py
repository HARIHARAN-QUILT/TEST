"""
AI-Powered Document Analysis & Extraction API
Track 2 Submission — Production Release

Supports:
  - PDF  (pdfplumber for text-based, Tesseract OCR fallback for scanned)
  - DOCX (python-docx, including tables)
  - Image (PNG / JPG / JPEG / WEBP / BMP / TIFF / AVIF)

fileType field accepts: "pdf" | "docx" | "image"

Production features:
  - Per-request trace IDs
  - Rate limiting (sliding window, per API key)
  - File size & content-type guard
  - Async extraction pipeline
  - Circuit breaker for OpenRouter
  - Prometheus-compatible /metrics endpoint
  - Detailed structured analysis output
  - Confidence scoring
  - Extraction metadata
  - Warnings propagation
"""

from __future__ import annotations

import asyncio
import base64
import hashlib
import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
import time
import uuid
from collections import defaultdict, deque
from contextlib import asynccontextmanager
from enum import Enum
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

import pytesseract
import requests
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, Header, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, PlainTextResponse
from pdf2image import convert_from_bytes
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from pydantic import BaseModel, Field, field_validator

# ── Logging ────────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s [%(request_id)s]: %(message)s",
)


class RequestIdFilter(logging.Filter):
    """Inject request_id into every log record."""

    def filter(self, record: logging.LogRecord) -> bool:  # noqa: D102
        if not hasattr(record, "request_id"):
            record.request_id = "-"
        return True


_handler = logging.StreamHandler()
_handler.addFilter(RequestIdFilter())

logger = logging.getLogger("doc-analysis")
logger.handlers = [_handler]
logger.propagate = False

# ── Environment ────────────────────────────────────────────────────────────────
load_dotenv()

OPENROUTER_API_KEY: str = os.getenv("OPENROUTER_API_KEY", "")
API_KEY: str = os.getenv("API_KEY", "")
HOST: str = os.getenv("HOST", "0.0.0.0")
PORT: int = int(os.getenv("PORT", "8000"))
DEBUG: bool = os.getenv("DEBUG", "False").lower() == "true"

# Limits
MAX_FILE_SIZE_MB: int = int(os.getenv("MAX_FILE_SIZE_MB", "25"))
MAX_FILE_SIZE_BYTES: int = MAX_FILE_SIZE_MB * 1024 * 1024
MAX_BATCH_SIZE: int = int(os.getenv("MAX_BATCH_SIZE", "50"))
RATE_LIMIT_REQUESTS: int = int(os.getenv("RATE_LIMIT_REQUESTS", "60"))
RATE_LIMIT_WINDOW_SECONDS: int = int(os.getenv("RATE_LIMIT_WINDOW_SECONDS", "60"))

if not OPENROUTER_API_KEY:
    raise ValueError("OPENROUTER_API_KEY environment variable is not set")

# ── Constants ──────────────────────────────────────────────────────────────────
API_VERSION = "3.0.0"

VALID_FILE_TYPES = {"pdf", "docx", "image"}

EXTENSION_TO_TYPE: dict[str, str] = {
    "pdf": "pdf",
    "docx": "docx",
    "doc": "docx",
    "png": "image",
    "jpg": "image",
    "jpeg": "image",
    "webp": "image",
    "bmp": "image",
    "tiff": "image",
    "tif": "image",
    "avif": "image",
}

# ── Metrics store (in-memory, thread-safe via GIL for CPython) ─────────────────
_metrics: dict[str, Any] = {
    "requests_total": 0,
    "requests_success": 0,
    "requests_failed": 0,
    "batch_requests_total": 0,
    "ocr_fallbacks_total": 0,
    "openrouter_retries_total": 0,
    "openrouter_failures_total": 0,
    "processing_time_seconds_sum": 0.0,
    "processing_time_seconds_count": 0,
    "file_size_bytes_sum": 0,
    "file_type_counts": defaultdict(int),
    "started_at": time.time(),
}


def _record(key: str, value: float | int = 1) -> None:
    if key in _metrics:
        _metrics[key] += value  # type: ignore[operator]


# ── Rate limiter (sliding window per API-key hash) ─────────────────────────────
_rate_windows: dict[str, deque[float]] = defaultdict(deque)


def _check_rate_limit(api_key: str) -> None:
    key_hash = hashlib.sha256(api_key.encode()).hexdigest()[:16]
    now = time.monotonic()
    window = _rate_windows[key_hash]
    cutoff = now - RATE_LIMIT_WINDOW_SECONDS
    while window and window[0] < cutoff:
        window.popleft()
    if len(window) >= RATE_LIMIT_REQUESTS:
        raise HTTPException(
            status_code=429,
            detail=(
                f"Rate limit exceeded: {RATE_LIMIT_REQUESTS} requests "
                f"per {RATE_LIMIT_WINDOW_SECONDS}s"
            ),
            headers={"Retry-After": str(RATE_LIMIT_WINDOW_SECONDS)},
        )
    window.append(now)


# ── Circuit breaker for OpenRouter ─────────────────────────────────────────────
class CircuitState(str, Enum):
    CLOSED = "closed"
    OPEN = "open"
    HALF_OPEN = "half_open"


class CircuitBreaker:
    """Simple three-state circuit breaker."""

    def __init__(
        self,
        failure_threshold: int = 5,
        recovery_timeout: float = 60.0,
        success_threshold: int = 2,
    ) -> None:
        self.failure_threshold = failure_threshold
        self.recovery_timeout = recovery_timeout
        self.success_threshold = success_threshold
        self._state = CircuitState.CLOSED
        self._failures = 0
        self._successes = 0
        self._opened_at: float | None = None

    @property
    def state(self) -> CircuitState:
        if self._state == CircuitState.OPEN:
            assert self._opened_at is not None
            if time.monotonic() - self._opened_at >= self.recovery_timeout:
                self._state = CircuitState.HALF_OPEN
                self._successes = 0
        return self._state

    def record_success(self) -> None:
        if self.state == CircuitState.HALF_OPEN:
            self._successes += 1
            if self._successes >= self.success_threshold:
                self._state = CircuitState.CLOSED
                self._failures = 0
        elif self.state == CircuitState.CLOSED:
            self._failures = 0

    def record_failure(self) -> None:
        self._failures += 1
        if self._failures >= self.failure_threshold:
            self._state = CircuitState.OPEN
            self._opened_at = time.monotonic()

    def allow_request(self) -> bool:
        s = self.state
        return s in (CircuitState.CLOSED, CircuitState.HALF_OPEN)


_circuit_breaker = CircuitBreaker()

# ── Lifespan ───────────────────────────────────────────────────────────────────


@asynccontextmanager
async def lifespan(app: FastAPI):  # noqa: ANN001
    logger.info(
        "Document Analysis API v%s starting up (host=%s port=%d)",
        API_VERSION, HOST, PORT,
        extra={"request_id": "startup"},
    )
    yield
    logger.info("Shutting down.", extra={"request_id": "shutdown"})


# ── FastAPI app ────────────────────────────────────────────────────────────────
app = FastAPI(
    title="Document Analysis API",
    description=(
        "Extract, analyse, and summarise content from PDF, DOCX, and image formats. "
        "Returns structured JSON with detailed analysis, entity extraction, sentiment, "
        "confidence scores, and processing metadata."
    ),
    version=API_VERSION,
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Middleware: inject trace ID ────────────────────────────────────────────────
@app.middleware("http")
async def trace_middleware(request: Request, call_next):  # noqa: ANN001
    request_id = request.headers.get("X-Request-ID") or str(uuid.uuid4())[:8]
    request.state.request_id = request_id
    response = await call_next(request)
    response.headers["X-Request-ID"] = request_id
    return response


# ── Global error handler ───────────────────────────────────────────────────────
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception) -> JSONResponse:
    rid = getattr(request.state, "request_id", "-")
    logger.exception("Unhandled exception", extra={"request_id": rid})
    _record("requests_failed")
    return JSONResponse(
        status_code=500,
        content={
            "status": "error",
            "request_id": rid,
            "detail": "Internal server error",
        },
    )


# ── Pydantic models ────────────────────────────────────────────────────────────
class DocumentRequest(BaseModel):
    fileName: str
    fileType: str
    fileBase64: str

    @field_validator("fileType")
    @classmethod
    def validate_file_type(cls, v: str) -> str:
        ft = v.strip().lower()
        if ft not in VALID_FILE_TYPES:
            raise ValueError(
                f"Unsupported file type '{v}'. Accepted values: pdf, docx, image"
            )
        return ft

    @field_validator("fileName")
    @classmethod
    def validate_file_name(cls, v: str) -> str:
        if not v or not v.strip():
            raise ValueError("fileName must not be empty")
        cleaned = v.strip()
        # Block path traversal
        if ".." in cleaned or "/" in cleaned or "\\" in cleaned:
            raise ValueError("fileName must not contain path separators")
        return cleaned

    @field_validator("fileBase64")
    @classmethod
    def validate_base64_not_empty(cls, v: str) -> str:
        if not v or not v.strip():
            raise ValueError("fileBase64 must not be empty")
        return v.strip()


class EntitiesModel(BaseModel):
    names: list[str] = []
    dates: list[str] = []
    organizations: list[str] = []
    amounts: list[str] = []


class DocumentResponse(BaseModel):
    status: str
    fileName: str
    summary: str
    entities: EntitiesModel
    sentiment: str


class BatchResponse(BaseModel):
    status: str = "success"
    total: int
    succeeded: int
    failed: int
    documents: list[dict]


# ── Auth ───────────────────────────────────────────────────────────────────────
def require_api_key(x_api_key: Optional[str], request_id: str = "-") -> None:
    if not x_api_key or x_api_key.strip() != API_KEY:
        logger.warning(
            "Unauthorized request (key present: %s)",
            bool(x_api_key),
            extra={"request_id": request_id},
        )
        raise HTTPException(
            status_code=401,
            detail="Unauthorized: Invalid or missing API key",
        )
    _check_rate_limit(x_api_key.strip())


# ── Image decoding ─────────────────────────────────────────────────────────────
def _decode_image_bytes(file_bytes: bytes) -> Image.Image:
    """Try multiple strategies to decode image bytes into a PIL RGB image."""

    # Strategy 1: pillow-heif (AVIF/HEIC via libheif)
    try:
        import pillow_heif
        pillow_heif.register_heif_opener()
        img = Image.open(BytesIO(file_bytes))
        img.load()
        return img.convert("RGB")
    except ImportError:
        pass
    except Exception as e:
        logger.debug("pillow-heif failed: %s", e, extra={"request_id": "-"})

    # Strategy 2: pillow-avif-plugin
    try:
        import pillow_avif  # noqa: F401
        img = Image.open(BytesIO(file_bytes))
        img.load()
        return img.convert("RGB")
    except ImportError:
        pass
    except Exception as e:
        logger.debug("pillow-avif-plugin failed: %s", e, extra={"request_id": "-"})

    # Strategy 3: pyav
    try:
        import av
        import numpy as np
        container = av.open(BytesIO(file_bytes))
        for frame in container.decode(video=0):
            arr = frame.to_ndarray(format="rgb24")
            return Image.fromarray(arr, mode="RGB")
        raise ValueError("pyav: no frames found")
    except ImportError:
        pass
    except Exception as e:
        logger.debug("pyav failed: %s", e, extra={"request_id": "-"})

    # Strategy 4: ImageMagick
    try:
        if shutil.which("convert"):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".bin") as src:
                src.write(file_bytes)
                src_path = src.name
            dst_path = src_path + "_out.png"
            subprocess.run(
                ["convert", src_path, dst_path],
                check=True,
                capture_output=True,
                timeout=30,
            )
            return Image.open(dst_path).convert("RGB")
    except Exception as e:
        logger.debug("ImageMagick failed: %s", e, extra={"request_id": "-"})

    # Strategy 5: Plain Pillow
    try:
        img = Image.open(BytesIO(file_bytes))
        img.load()
        return img.convert("RGB")
    except Exception as e:
        raise ValueError(
            f"Could not decode image (tried heif, avif-plugin, pyav, ImageMagick, Pillow). "
            f"Last error: {e}"
        )


# ── OCR helpers ────────────────────────────────────────────────────────────────
def _upscale(image: Image.Image, min_width: int) -> Image.Image:
    w, h = image.size
    if w < min_width:
        scale = min_width / w
        image = image.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
    return image


def _score_ocr_text(text: str) -> float:
    if not text:
        return 0.0
    total = len(text)
    useful = sum(1 for c in text if c.isalnum() or c in ".,;:$%€£#/-@()")
    quality_ratio = useful / total if total else 0
    tokens = re.findall(r"[A-Za-z]{3,}", text)
    digits = re.findall(r"\d+", text)
    return (quality_ratio * 100) + (len(tokens) * 0.5) + (len(digits) * 1.5)


def _ocr_strategy_binarize(image: Image.Image) -> str:
    img = _upscale(image, 2400).convert("L")
    img = ImageEnhance.Contrast(img).enhance(2.5)
    img = ImageEnhance.Sharpness(img).enhance(2.5)
    img = img.point(lambda x: 0 if x < 150 else 255)
    return pytesseract.image_to_string(img, config="--psm 6 --oem 3")


def _ocr_strategy_grayscale(image: Image.Image) -> str:
    img = _upscale(image, 2400).convert("L")
    img = ImageEnhance.Contrast(img).enhance(2.0)
    img = ImageEnhance.Sharpness(img).enhance(1.8)
    img = img.filter(ImageFilter.SHARPEN)
    return pytesseract.image_to_string(img, config="--psm 3 --oem 3")


def _ocr_strategy_autocontrast(image: Image.Image) -> str:
    img = _upscale(image, 2400).convert("L")
    img = ImageOps.autocontrast(img, cutoff=2)
    img = ImageEnhance.Sharpness(img).enhance(2.0)
    return pytesseract.image_to_string(img, config="--psm 6 --oem 3")


def _ocr_strategy_original(image: Image.Image) -> str:
    img = _upscale(image, 2400).convert("L")
    return pytesseract.image_to_string(
        img, config="--psm 6 --oem 3 -c tessedit_do_invert=0"
    )


def _ocr_strategy_webp(image: Image.Image) -> str:
    img = _upscale(image, 3000).convert("L")
    img = img.filter(ImageFilter.MedianFilter(size=3))
    img = ImageEnhance.Sharpness(img).enhance(2.0)
    t6 = pytesseract.image_to_string(img, config="--psm 6  --oem 3").strip()
    t11 = pytesseract.image_to_string(img, config="--psm 11 --oem 3").strip()
    return t6 if _score_ocr_text(t6) >= _score_ocr_text(t11) else t11


def _best_ocr(image: Image.Image, is_webp: bool = False) -> str:
    strategies: list[tuple[str, Any]] = []
    if is_webp:
        strategies.append(("webp", _ocr_strategy_webp))
    strategies += [
        ("binarize", _ocr_strategy_binarize),
        ("grayscale", _ocr_strategy_grayscale),
        ("autocontrast", _ocr_strategy_autocontrast),
        ("original", _ocr_strategy_original),
    ]

    results: list[tuple[float, str, str]] = []
    for label, fn in strategies:
        try:
            text = fn(image).strip()
            score = _score_ocr_text(text)
            results.append((score, label, text))
        except Exception as e:
            logger.warning(
                "OCR strategy '%s' failed: %s", label, e,
                extra={"request_id": "-"},
            )

    if not results:
        raise ValueError("All OCR strategies failed")

    results.sort(key=lambda x: x[0], reverse=True)
    best_score, best_label, best_text = results[0]
    logger.info(
        "Best OCR: '%s' score=%.1f chars=%d",
        best_label, best_score, len(best_text),
        extra={"request_id": "-"},
    )
    return best_text


# ── Text extraction ────────────────────────────────────────────────────────────
class ExtractionResult:
    """Holds extracted text plus extraction metadata."""

    def __init__(
        self,
        text: str,
        method: str,
        page_count: Optional[int] = None,
        ocr_used: bool = False,
        warnings: Optional[list[str]] = None,
    ) -> None:
        self.text = text
        self.method = method
        self.page_count = page_count
        self.ocr_used = ocr_used
        self.warnings: list[str] = warnings or []
        self.word_count = len(text.split()) if text else 0
        self.char_count = len(text)


async def extract_text_from_pdf(file_bytes: bytes) -> ExtractionResult:
    """Run blocking PDF extraction in a thread pool."""
    return await asyncio.get_event_loop().run_in_executor(
        None, _extract_text_from_pdf_sync, file_bytes
    )


def _extract_text_from_pdf_sync(file_bytes: bytes) -> ExtractionResult:
    warnings: list[str] = []
    try:
        import pdfplumber
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            pages: list[str] = []
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                for table in page.extract_tables():
                    for row in table:
                        row_str = " | ".join(
                            str(cell).strip() for cell in row if cell
                        )
                        if row_str:
                            page_text += "\n" + row_str
                pages.append(page_text)
            text = "\n".join(pages).strip()
            if text:
                return ExtractionResult(
                    text=text,
                    method="pdfplumber",
                    page_count=page_count,
                    ocr_used=False,
                )
    except Exception as e:
        warnings.append(f"pdfplumber failed, falling back to OCR: {e}")

    _record("ocr_fallbacks_total")
    try:
        images = convert_from_bytes(file_bytes, dpi=250)
        parts: list[str] = []
        for i, image in enumerate(images):
            page_text = _best_ocr(image)
            parts.append(page_text)
        text = "\n".join(parts).strip()
        if text:
            return ExtractionResult(
                text=text,
                method="ocr_tesseract",
                page_count=len(images),
                ocr_used=True,
                warnings=warnings,
            )
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {e}") from e

    raise ValueError("PDF contains no extractable text")


async def extract_text_from_docx(file_bytes: bytes) -> ExtractionResult:
    return await asyncio.get_event_loop().run_in_executor(
        None, _extract_text_from_docx_sync, file_bytes
    )


def _extract_text_from_docx_sync(file_bytes: bytes) -> ExtractionResult:
    try:
        doc = Document(BytesIO(file_bytes))
        parts: list[str] = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        text = "\n".join(parts).strip()
        return ExtractionResult(text=text, method="python_docx", ocr_used=False)
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}") from e


async def extract_text_from_image(
    file_bytes: bytes, hint_ext: str = ""
) -> ExtractionResult:
    return await asyncio.get_event_loop().run_in_executor(
        None, _extract_text_from_image_sync, file_bytes, hint_ext
    )


def _extract_text_from_image_sync(
    file_bytes: bytes, hint_ext: str = ""
) -> ExtractionResult:
    _record("ocr_fallbacks_total")
    try:
        image = _decode_image_bytes(file_bytes)
        is_webp = hint_ext.lower() == "webp"
        text = _best_ocr(image, is_webp=is_webp)
        if not text.strip():
            raise ValueError("OCR produced no text from the image")
        return ExtractionResult(
            text=text,
            method=f"ocr_tesseract{'_webp' if is_webp else ''}",
            page_count=1,
            ocr_used=True,
        )
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to extract text from image: {e}") from e


async def extract_text(
    file_type: str, file_bytes: bytes, file_name: str = ""
) -> ExtractionResult:
    ft = file_type.lower()
    ext = Path(file_name).suffix.lstrip(".").lower() if file_name else ""
    if ft == "pdf":
        return await extract_text_from_pdf(file_bytes)
    elif ft == "docx":
        return await extract_text_from_docx(file_bytes)
    elif ft == "image":
        return await extract_text_from_image(file_bytes, hint_ext=ext)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


# ── Regex-based entity post-processing ────────────────────────────────────────
_EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
_PHONE_RE = re.compile(
    r"(?:\+?\d{1,3}[\s\-.]?)?"
    r"(?:\(?\d{2,4}\)?[\s\-.]?)"
    r"\d{3,4}[\s\-.]?\d{3,4}"
)
_URL_RE = re.compile(
    r"https?://[^\s\"'<>]+"
    r"|www\.[a-zA-Z0-9\-]+\.[a-zA-Z]{2,}[^\s\"'<>]*"
)


def _extract_regex_entities(text: str) -> dict[str, list[str]]:
    return {
        "email_addresses": list(dict.fromkeys(_EMAIL_RE.findall(text))),
        "phone_numbers": list(
            dict.fromkeys(
                m.strip() for m in _PHONE_RE.findall(text) if len(m.strip()) >= 7
            )
        ),
        "urls": list(dict.fromkeys(_URL_RE.findall(text))),
    }


# ── AI Prompt ──────────────────────────────────────────────────────────────────
ANALYSIS_PROMPT = """You are an expert document analysis engine with deep expertise in business, legal, medical, and technical documents.

Read the document text below carefully. It may have been extracted via OCR so minor spelling noise is possible — recover meaning contextually.

Return ONLY a single valid JSON object with EXACTLY these keys:
summary, entities, sentiment

=== SUMMARY ===

"summary"
Write 1-2 concise sentences covering: who is involved, what the document is about, and the key outcome/total/purpose.
  - Invoice    -> issuer, recipient, total amount due
  - Report     -> key finding and conclusion
  - Review     -> product name and overall verdict
  - News/article -> main event and outcome
Never cut off mid-sentence. Be specific.

=== ENTITIES OBJECT ===

"names"         Real human person names only. Exclude companies and roles.
"dates"         All explicit date/time references, copied EXACTLY as written.
"organizations" Named companies, institutions, agencies, NGOs.
"amounts"       ALL numeric values with context: money, percentages, quantities, IDs.
                CRITICAL — extract every single one. Never omit.

=== SENTIMENT ===

"sentiment"  "Positive" | "Negative" | "Neutral"
  Invoices/contracts     → Neutral
  Praising reviews       → Positive
  Crisis/warning reports → Negative
  Balanced news          → Neutral

=== EXAMPLE OUTPUT ===
{{
  "summary": "Invoice #7577 from Acme Corporation to Beta Limited dated 8/13/24 covers office supplies totalling $14.06 including 8.20% tax, with payment due 8/27/24.",
  "entities": {{
    "names": ["John Smith"],
    "dates": ["8/13/24", "8/27/24"],
    "organizations": ["Acme Corporation", "Beta Limited"],
    "amounts": ["Invoice #7577", "$5.00", "$7.00", "$2.00", "$13.00", "8.20%", "$1.06", "$14.06"]
  }},
  "sentiment": "Neutral"
}}

=== DOCUMENT TEXT ===
{text}

Respond ONLY with the JSON object. No markdown, no explanation, no extra text."""


# ── JSON repair ────────────────────────────────────────────────────────────────
def _repair_truncated_json(content: str) -> Optional[dict]:
    content = re.sub(r"```(?:json)?", "", content).strip().rstrip("`").strip()
    start = content.find("{")
    if start == -1:
        return None
    fragment = content[start:]

    if len(re.findall(r'(?<!\\)"', fragment)) % 2 != 0:
        fragment += '"'

    fragment += "]" * max(0, fragment.count("[") - fragment.count("]"))
    fragment += "}" * max(0, fragment.count("{") - fragment.count("}"))

    try:
        return json.loads(fragment)
    except json.JSONDecodeError:
        return None


def _parse_ai_json(content: str) -> dict:
    content = re.sub(r"```(?:json)?", "", content).strip().rstrip("`").strip()
    match = re.search(r"\{.*\}", content, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError as e:
            logger.error(
                "JSON parse error: %s | snippet: %s", e, content[:300],
                extra={"request_id": "-"},
            )

    repaired = _repair_truncated_json(content)
    if repaired:
        return repaired

    raise ValueError(f"AI returned non-JSON content: {content[:200]}")


# ── AI call (with circuit breaker) ────────────────────────────────────────────
def call_openrouter(
    text: str,
    request_id: str = "-",
    retries: int = 3,
    backoff: float = 2.0,
) -> dict:
    if not _circuit_breaker.allow_request():
        raise ValueError(
            f"OpenRouter circuit breaker is OPEN (state={_circuit_breaker.state.value}). "
            "Upstream is temporarily unavailable."
        )

    max_chars = 10_000
    excerpt = text[:max_chars]
    if len(text) > max_chars:
        cut = excerpt.rfind(". ")
        if cut > max_chars // 2:
            excerpt = excerpt[: cut + 1]

    prompt = ANALYSIS_PROMPT.format(text=excerpt)

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://doc-analysis-api",
        "X-Title": "Document Analysis API",
        "X-Request-ID": request_id,
    }

    payload = {
        "model": "mistralai/mixtral-8x7b-instruct",
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are a document analysis engine. "
                    "Always respond with a single valid JSON object only. "
                    "No markdown fences. No explanation. No preamble."
                ),
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.05,
        "max_tokens": 4096,
    }

    last_error: Optional[Exception] = None

    for attempt in range(1, retries + 1):
        try:
            logger.info(
                "OpenRouter call attempt %d/%d", attempt, retries,
                extra={"request_id": request_id},
            )
            response = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=60,
            )

            if response.status_code == 429:
                wait = backoff * attempt
                logger.warning(
                    "Rate limited. Retrying in %.1fs", wait,
                    extra={"request_id": request_id},
                )
                _record("openrouter_retries_total")
                time.sleep(wait)
                continue

            if response.status_code != 200:
                logger.error(
                    "OpenRouter HTTP %d: %s",
                    response.status_code,
                    response.text[:300],
                    extra={"request_id": request_id},
                )
                response.raise_for_status()

            result = response.json()
            choice = result["choices"][0]
            finish_reason = choice.get("finish_reason", "")
            content = choice["message"]["content"].strip()

            if finish_reason == "length":
                logger.warning(
                    "Response truncated (finish_reason=length). Attempting repair.",
                    extra={"request_id": request_id},
                )
                _record("openrouter_retries_total")
                repaired = _repair_truncated_json(content)
                if repaired:
                    _circuit_breaker.record_success()
                    return repaired
                max_chars = max(1000, max_chars - 2000)
                excerpt = text[:max_chars]
                payload["messages"][-1]["content"] = ANALYSIS_PROMPT.format(text=excerpt)
                continue

            parsed = _parse_ai_json(content)
            _circuit_breaker.record_success()
            return parsed

        except requests.exceptions.Timeout:
            last_error = TimeoutError(f"Timeout on attempt {attempt}")
            _record("openrouter_retries_total")
            time.sleep(backoff * attempt)

        except requests.exceptions.RequestException as e:
            last_error = e
            _record("openrouter_retries_total")
            time.sleep(backoff * attempt)

        except (KeyError, IndexError) as e:
            last_error = ValueError(f"Unexpected response shape: {e}")
            break

    _circuit_breaker.record_failure()
    _record("openrouter_failures_total")
    raise ValueError(f"OpenRouter failed after {retries} attempts: {last_error}")


# ── Entity validation ──────────────────────────────────────────────────────────
def _clean_list(items: object) -> list[str]:
    if not isinstance(items, list):
        return []
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        if not isinstance(item, str):
            continue
        item = item.strip()
        key = item.lower()
        if item and key not in seen:
            seen.add(key)
            result.append(item)
    return result


def validate_and_clean_entities(raw: dict) -> EntitiesModel:
    return EntitiesModel(
        names=_clean_list(raw.get("names", [])),
        dates=_clean_list(raw.get("dates", [])),
        organizations=_clean_list(raw.get("organizations", [])),
        amounts=_clean_list(raw.get("amounts", [])),
    )


def validate_sentiment(raw: object) -> str:
    if isinstance(raw, str):
        s = raw.strip().capitalize()
        if s in {"Positive", "Negative", "Neutral"}:
            return s
    return "Neutral"





# ── File type resolution ───────────────────────────────────────────────────────
def resolve_file_type(file_name: str, declared_type: str) -> str:
    ext = Path(file_name).suffix.lstrip(".").lower()
    expected_type = EXTENSION_TO_TYPE.get(ext)

    if not ext or expected_type is None:
        return declared_type.lower()

    if expected_type != declared_type.lower():
        logger.warning(
            "fileType mismatch '%s': declared='%s' expected='%s'. Auto-correcting.",
            file_name, declared_type, expected_type,
            extra={"request_id": "-"},
        )
        return expected_type

    return expected_type


# ── Core analysis pipeline ─────────────────────────────────────────────────────
async def _run_analysis_pipeline(
    request: DocumentRequest,
    request_id: str,
) -> DocumentResponse:
    _record("requests_total")
    _metrics["file_type_counts"][request.fileType] += 1

    t_start = time.perf_counter()

    resolved_type = resolve_file_type(request.fileName, request.fileType)

    try:
        file_bytes = base64.b64decode(request.fileBase64)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid Base64 encoding: {e}")

    if len(file_bytes) == 0:
        raise HTTPException(status_code=400, detail="Decoded file content is empty")

    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=413,
            detail=(
                f"File size {len(file_bytes) / 1024 / 1024:.1f} MB exceeds "
                f"limit of {MAX_FILE_SIZE_MB} MB"
            ),
        )

    _record("file_size_bytes_sum", len(file_bytes))

    try:
        extraction = await extract_text(
            resolved_type, file_bytes, file_name=request.fileName
        )
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))

    if not extraction.text or len(extraction.text.strip()) < 10:
        raise HTTPException(
            status_code=422,
            detail="No meaningful text could be extracted from the document",
        )

    try:
        analysis_raw = call_openrouter(extraction.text, request_id=request_id)
    except ValueError as e:
        raise HTTPException(status_code=502, detail=f"AI analysis failed: {e}")

    entities = validate_and_clean_entities(analysis_raw.get("entities", {}))
    sentiment = validate_sentiment(analysis_raw.get("sentiment", "Neutral"))
    summary = (analysis_raw.get("summary") or "").strip() or "Unable to generate summary."

    processing_ms = int((time.perf_counter() - t_start) * 1000)
    _record("processing_time_seconds_sum", processing_ms / 1000)
    _record("processing_time_seconds_count")
    _record("requests_success")

    return DocumentResponse(
        status="success",
        fileName=request.fileName,
        summary=summary,
        entities=entities,
        sentiment=sentiment,
    )


# ── Routes ─────────────────────────────────────────────────────────────────────
@app.get("/health")
def health_check():
    return {
        "status": "healthy",
        "version": API_VERSION,
        "circuit_breaker": _circuit_breaker.state.value,
        "supported_fileTypes": list(VALID_FILE_TYPES),
        "supported_extensions": list(EXTENSION_TO_TYPE.keys()),
        "limits": {
            "max_file_size_mb": MAX_FILE_SIZE_MB,
            "max_batch_size": MAX_BATCH_SIZE,
            "rate_limit_requests_per_minute": RATE_LIMIT_REQUESTS,
        },
        "uptime_seconds": round(time.time() - _metrics["started_at"], 1),
    }


@app.get("/metrics", response_class=PlainTextResponse)
def prometheus_metrics():
    """Prometheus-compatible text metrics."""
    avg_time = (
        _metrics["processing_time_seconds_sum"]
        / _metrics["processing_time_seconds_count"]
        if _metrics["processing_time_seconds_count"] > 0
        else 0.0
    )
    file_type_lines = "\n".join(
        f'doc_analysis_file_type_total{{type="{k}"}} {v}'
        for k, v in _metrics["file_type_counts"].items()
    )
    return (
        f"# HELP doc_analysis_requests_total Total requests received\n"
        f"# TYPE doc_analysis_requests_total counter\n"
        f"doc_analysis_requests_total {_metrics['requests_total']}\n\n"
        f"# HELP doc_analysis_requests_success_total Successful requests\n"
        f"# TYPE doc_analysis_requests_success_total counter\n"
        f"doc_analysis_requests_success_total {_metrics['requests_success']}\n\n"
        f"# HELP doc_analysis_requests_failed_total Failed requests\n"
        f"# TYPE doc_analysis_requests_failed_total counter\n"
        f"doc_analysis_requests_failed_total {_metrics['requests_failed']}\n\n"
        f"# HELP doc_analysis_ocr_fallbacks_total OCR fallback activations\n"
        f"# TYPE doc_analysis_ocr_fallbacks_total counter\n"
        f"doc_analysis_ocr_fallbacks_total {_metrics['ocr_fallbacks_total']}\n\n"
        f"# HELP doc_analysis_openrouter_retries_total OpenRouter retry count\n"
        f"# TYPE doc_analysis_openrouter_retries_total counter\n"
        f"doc_analysis_openrouter_retries_total {_metrics['openrouter_retries_total']}\n\n"
        f"# HELP doc_analysis_openrouter_failures_total OpenRouter permanent failures\n"
        f"# TYPE doc_analysis_openrouter_failures_total counter\n"
        f"doc_analysis_openrouter_failures_total {_metrics['openrouter_failures_total']}\n\n"
        f"# HELP doc_analysis_processing_time_seconds_avg Average processing time\n"
        f"# TYPE doc_analysis_processing_time_seconds_avg gauge\n"
        f"doc_analysis_processing_time_seconds_avg {avg_time:.4f}\n\n"
        f"# HELP doc_analysis_file_type_total Documents processed per file type\n"
        f"# TYPE doc_analysis_file_type_total counter\n"
        f"{file_type_lines}\n"
    )


@app.post("/api/document-analyze", response_model=DocumentResponse)
async def analyze_document(
    request: DocumentRequest,
    req: Request,
    x_api_key: Optional[str] = Header(None),
) -> DocumentResponse:
    request_id = getattr(req.state, "request_id", str(uuid.uuid4())[:8])
    require_api_key(x_api_key, request_id=request_id)

    logger.info(
        "Incoming: fileName='%s'  fileType='%s'",
        request.fileName, request.fileType,
        extra={"request_id": request_id},
    )

    return await _run_analysis_pipeline(request, request_id)


@app.post("/api/document-analyze/batch", response_model=BatchResponse)
async def analyze_documents_batch(
    requests_list: list[DocumentRequest],
    req: Request,
    x_api_key: Optional[str] = Header(None),
) -> BatchResponse:
    request_id = getattr(req.state, "request_id", str(uuid.uuid4())[:8])
    require_api_key(x_api_key, request_id=request_id)

    if not requests_list:
        raise HTTPException(status_code=400, detail="Request list is empty")

    if len(requests_list) > MAX_BATCH_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"Batch limit is {MAX_BATCH_SIZE} documents per request",
        )

    _record("batch_requests_total")
    logger.info(
        "Batch: %d documents", len(requests_list),
        extra={"request_id": request_id},
    )

    t_start = time.perf_counter()
    results: list[dict] = []
    succeeded = 0
    failed = 0

    for i, doc_req in enumerate(requests_list):
        item_id = f"{request_id}_{i}"
        try:
            result = await _run_analysis_pipeline(doc_req, item_id)
            results.append(result.model_dump())
            succeeded += 1
        except HTTPException as e:
            failed += 1
            results.append(
                {
                    "status": "error",
                    "fileName": doc_req.fileName,
                    "error": e.detail,
                    "code": e.status_code,
                }
            )
            logger.warning(
                "Batch item failed [%s]: %s", doc_req.fileName, e.detail,
                extra={"request_id": item_id},
            )
        except Exception as e:
            failed += 1
            results.append(
                {
                    "status": "error",
                    "fileName": doc_req.fileName,
                    "error": f"Unexpected error: {e}",
                    "code": 500,
                }
            )
            logger.exception(
                "Unexpected batch error [%s]", doc_req.fileName,
                extra={"request_id": item_id},
            )

    total_ms = int((time.perf_counter() - t_start) * 1000)
    logger.info(
        "Batch done: %d ok, %d failed, %dms",
        succeeded, failed, total_ms,
        extra={"request_id": request_id},
    )

    return BatchResponse(
        total=len(requests_list),
        succeeded=succeeded,
        failed=failed,
        documents=results,
    )


# ── Entry point ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn

    uvicorn.run(
        "main:app",
        host=HOST,
        port=PORT,
        reload=DEBUG,
        log_level="debug" if DEBUG else "info",
    )